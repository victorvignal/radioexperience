from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import re
import win32com.client

src = Path(r'C:\Users\vigna\.openclaw\workspace\RadioeXperienceRAG\catalog\Stack_MVP_RAG_Chatbot_Site_Admin_V3.docx')
out_docx = Path(r'C:\Users\vigna\.openclaw\workspace\RadioeXperienceRAG\catalog\Stack_MVP_RAG_Chatbot_Site_Admin_V3_Organizado.docx')
out_pdf = Path(r'C:\Users\vigna\.openclaw\workspace\RadioeXperienceRAG\catalog\Stack_MVP_RAG_Chatbot_Site_Admin_V3_Organizado.pdf')

doc = Document(str(src))
styles = doc.styles
styles['Normal'].font.name = 'Calibri'
styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
styles['Normal'].font.size = Pt(11)

# Light cleanup: normalize spacing and headings punctuation
for p in doc.paragraphs:
    txt = p.text
    txt = re.sub(r'\s+', ' ', txt).strip()
    if not txt:
        continue
    if p.style.name.startswith('Heading'):
        txt = txt.replace(' :', ':')
    # rewrite paragraph cleanly
    if p.text != txt:
        for r in p.runs:
            r.text = ''
        if p.runs:
            p.runs[0].text = txt
        else:
            p.add_run(txt)

# Add small footer note section if not present
last = doc.paragraphs[-1].text if doc.paragraphs else ''
if 'Observação final' not in last:
    doc.add_heading('Observação final', level=1)
    doc.add_paragraph('Este documento foi reorganizado para leitura mais limpa. Preços de APIs e plataformas podem mudar; vale conferir os links oficiais no momento de contratar.')

doc.save(str(out_docx))

# Export to PDF via Word COM
word = win32com.client.Dispatch('Word.Application')
word.Visible = False
wdFormatPDF = 17
wb = word.Documents.Open(str(out_docx))
wb.SaveAs(str(out_pdf), FileFormat=wdFormatPDF)
wb.Close(False)
word.Quit()

print(out_docx)
print(out_pdf)
