from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

out_path = r'C:\Users\vigna\.openclaw\workspace\RadioeXperienceRAG\catalog\Stack_MVP_RAG_Chatbot_Site_Admin_V2.docx'

doc = Document()
styles = doc.styles
styles['Normal'].font.name = 'Calibri'
styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
styles['Normal'].font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Stack recomendado para o MVP — V2\nRAG + Chatbot + Site + Login + Painel Admin')
r.bold = True
r.font.size = Pt(18)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Versão mais executiva, organizada e comparativa — março/2026')
r.italic = True
r.font.size = Pt(10)


def H(text, level=1):
    doc.add_heading(text, level=level)

def P(text):
    doc.add_paragraph(text)

def B(text, level=0):
    doc.add_paragraph(text, style='List Bullet' if level == 0 else 'List Bullet 2')

H('1. Objetivo', 1)
P('Definir uma stack enxuta, escalável e pragmática para lançar um MVP com base documental, busca semântica, chatbot com RAG, autenticação de usuários, site institucional e painel administrativo.')

H('2. Stack recomendada (decisão principal)', 1)
B('Frontend e painel: Next.js')
B('Autenticação: Supabase Auth')
B('Banco do app: Postgres')
B('Backend de IA / RAG: FastAPI')
B('Banco vetorial: Qdrant')
B('Storage de arquivos: Cloudflare R2')
B('Deploy frontend: Vercel')
B('Deploy backend Python: Railway')

P('Essa combinação separa bem o app web do núcleo de IA. O site e o painel ficam produtivos em Next.js; a parte de ingestão, parsing, embeddings e retrieval fica em Python, onde o ecossistema de RAG é mais natural.')

H('3. Arquitetura em blocos', 1)
P('Fluxo sugerido:')
B('Usuário entra no site ou área logada (Next.js)')
B('Login e sessão via Supabase Auth')
B('Frontend chama API do backend FastAPI para chat, busca e ingestão')
B('FastAPI consulta Postgres para dados do app e Qdrant para recuperação semântica')
B('Arquivos PDF, imagens extraídas e artefatos ficam no Cloudflare R2')
B('Frontend deployado na Vercel; backend Python na Railway')

H('4. Ferramenta por ferramenta', 1)

H('4.1 Next.js', 2)
P('Uso no projeto: site institucional, páginas públicas, área logada, painel administrativo e interface do chat.')
B('Ponto forte: produtividade alta, excelente DX e ecossistema maduro.')
B('Alternativas: Remix, Nuxt, Astro.')

H('4.2 Supabase Auth', 2)
P('Uso no projeto: login, sessão, recuperação de senha, proteção de rotas e gestão inicial de usuários.')
B('Ponto forte: reduz muito o tempo de construção do MVP.')
B('Alternativas: Auth.js, Clerk, Auth0.')

H('4.3 Postgres', 2)
P('Uso no projeto: usuários, roles, sessões de chat, documentos cadastrados, logs operacionais, jobs e configurações.')
B('Ponto forte: banco transacional robusto e previsível.')
B('Alternativas: MySQL, Neon, Postgres do próprio Supabase.')

H('4.4 FastAPI', 2)
P('Uso no projeto: upload, ingestão, chunking, embeddings, retrieval, endpoints do chat e tarefas assíncronas.')
B('Ponto forte: encaixa melhor com parsing, NLP e bibliotecas de IA em Python.')
B('Alternativas: Flask, Django Ninja, backend full-stack em Next.js com workers Python.')

H('4.5 Qdrant', 2)
P('Uso no projeto: indexação vetorial dos chunks, filtros por metadata e busca semântica no fluxo de RAG.')
B('Ponto forte: excelente relação simplicidade/desempenho para RAG.')
B('Alternativas: pgvector, Pinecone, Weaviate, Milvus.')

H('4.6 Cloudflare R2', 2)
P('Uso no projeto: PDFs, imagens, artefatos do pipeline, arquivos de apoio e eventualmente exportações.')
B('Ponto forte: storage barato com egress muito amigável.')
B('Alternativas: AWS S3, Supabase Storage, Vercel Blob.')

H('4.7 Vercel', 2)
P('Uso no projeto: deploy do frontend e da parte web em Next.js.')
B('Ponto forte: melhor experiência para Next.js.')
B('Alternativas: Netlify, Cloudflare Pages, VPS próprio.')

H('4.8 Railway', 2)
P('Uso no projeto: deploy do FastAPI, workers, jobs e serviços Python auxiliares.')
B('Ponto forte: simplicidade operacional para MVP.')
B('Alternativas: Render, Fly.io, VPS próprio.')

H('5. Comparativo das principais escolhas', 1)

t = doc.add_table(rows=1, cols=5)
t.style = 'Table Grid'
headers = ['Camada', 'Escolha principal', 'Por que ela ganhou', 'Alternativa forte', 'Quando escolher a alternativa']
for i, h in enumerate(headers):
    t.rows[0].cells[i].text = h
rows = [
    ['Frontend', 'Next.js', 'Melhor encaixe para site + app + admin', 'Remix', 'Se quiser uma abordagem mais minimalista de rotas/actions'],
    ['Auth', 'Supabase Auth', 'Mais rápido para MVP', 'Auth.js', 'Se quiser controle total e menos dependência de vendor'],
    ['Backend IA', 'FastAPI', 'Python conversa melhor com RAG', 'Next.js API + workers Python', 'Se quiser centralizar a camada web em TS'],
    ['Banco vetorial', 'Qdrant', 'Boa experiência para RAG e filtros', 'pgvector', 'Se quiser reduzir número de serviços'],
    ['Storage', 'Cloudflare R2', 'Barato e amigável para arquivos', 'S3', 'Se já estiver todo na AWS'],
    ['Deploy backend', 'Railway', 'Deploy simples para MVP', 'Render', 'Se preferir um modelo mais tradicional de serviços'],
]
for row in rows:
    c = t.add_row().cells
    for i, val in enumerate(row):
        c[i].text = val

H('6. Preços pesquisados (referência inicial)', 1)
P('Os valores abaixo são referências públicas pesquisadas em março/2026. Preço real depende de uso e região.')

t2 = doc.add_table(rows=1, cols=4)
t2.style = 'Table Grid'
headers = ['Ferramenta', 'Preço', 'Como entra no projeto', 'Observação']
for i, h in enumerate(headers):
    t2.rows[0].cells[i].text = h
pricing_rows = [
    ['OpenAI GPT-5.4', 'US$2.50/1M input, US$15/1M output', 'Respostas do chat / raciocínio', 'Modelo forte, mas não o mais barato'],
    ['OpenAI GPT-5.4 mini', 'US$0.75/1M input, US$4.50/1M output', 'Chat mais econômico / operações frequentes', 'Boa opção de custo-benefício'],
    ['OpenAI text-embedding-3-small', '≈ US$0.02 / 1M tokens', 'Embeddings para indexação RAG', 'Confirmar no painel oficial antes da contratação'],
    ['Qdrant Cloud', 'Free tier + Standard por uso', 'Banco vetorial', 'Cobra compute, RAM, storage e backups'],
    ['Supabase', 'Free; Pro a partir de US$25/mês', 'Auth + Postgres opcional + recursos do app', 'Muito bom para MVP'],
    ['Vercel', 'Hobby grátis; Pro US$20/usuário/mês + uso', 'Deploy do frontend', 'Excelente DX com Next.js'],
    ['Railway', 'US$5 Hobby mínimo / US$20 Pro mínimo + uso', 'Deploy do backend Python', 'Cobrança por CPU, RAM, volume e egress'],
    ['Cloudflare R2', 'US$0.015/GB-mês + operações', 'Storage dos arquivos', 'Tem free tier útil'],
    ['Auth.js', 'Grátis', 'Alternativa de auth', 'Open source; infra paga à parte'],
]
for row in pricing_rows:
    c = t2.add_row().cells
    for i, val in enumerate(row):
        c[i].text = val

H('7. Estimativa de custo por fase', 1)
B('Protótipo técnico: US$0–50/mês + uso de API')
B('MVP fechado para testes: US$50–150/mês + uso de OpenAI')
B('MVP rodando com usuários reais: US$150–400/mês, dependendo de tráfego, volume de PDFs e uso de chat')
P('A maior variável tende a ser o consumo de LLM e embeddings, não a hospedagem em si.')

H('8. Ordem de implementação recomendada', 1)
B('Fase 1: estruturar base documental e pipeline de ingestão')
B('Fase 2: subir backend FastAPI com endpoints de upload, ingest e search')
B('Fase 3: conectar Qdrant e embeddings')
B('Fase 4: criar frontend Next.js com login e chat')
B('Fase 5: montar painel admin para documentos, usuários e sessões')
B('Fase 6: observabilidade, logs, rate limit e hardening')

H('9. Alternativas de arquitetura', 1)

H('9.1 Arquitetura mais simples', 2)
B('Next.js + Supabase + Qdrant + R2 + scripts Python')
B('Boa se quiser reduzir número de serviços no começo')

H('9.2 Arquitetura mais coesa para IA', 2)
B('Next.js + FastAPI + Postgres + Qdrant + R2')
B('Melhor separação entre produto web e motor de RAG')

H('9.3 Arquitetura minimalista com menos peças', 2)
B('Next.js + Postgres + pgvector + Auth.js')
B('Menos componentes, mas menos especializado para pipeline pesado')

H('10. Recomendação final', 1)
P('Se o objetivo é lançar rápido sem criar dívida técnica desnecessária, a melhor escolha para este projeto hoje é: Next.js + Supabase Auth + Postgres + FastAPI + Qdrant + Cloudflare R2 + Vercel + Railway.')
P('É uma stack prática, relativamente barata para um MVP, e boa o bastante para crescer depois.')

H('11. Fontes', 1)
B('OpenAI pricing: https://openai.com/api/pricing/')
B('Qdrant pricing: https://qdrant.tech/pricing/')
B('Supabase pricing: https://supabase.com/pricing')
B('Vercel pricing: https://vercel.com/pricing')
B('Railway pricing: https://railway.com/pricing')
B('Cloudflare R2 pricing: https://developers.cloudflare.com/r2/pricing/')
B('Auth.js: https://authjs.dev/')

doc.save(out_path)
print(out_path)
