from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

out_path = r'C:\Users\vigna\.openclaw\workspace\RadioeXperienceRAG\catalog\Stack_MVP_RAG_Chatbot_Site_Admin_V3.docx'

doc = Document()
styles = doc.styles
styles['Normal'].font.name = 'Calibri'
styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
styles['Normal'].font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Stack recomendado para o MVP — V3\nRAG + Chatbot + Site + Login + Painel Admin')
r.bold = True
r.font.size = Pt(18)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Organizado por serviços necessários, com custos e alternativas — março/2026')
r.italic = True
r.font.size = Pt(10)


def H(text, level=1):
    doc.add_heading(text, level=level)

def P(text):
    doc.add_paragraph(text)

def B(text, level=0):
    doc.add_paragraph(text, style='List Bullet' if level == 0 else 'List Bullet 2')

H('1. Visão geral', 1)
P('A ideia deste documento é separar o projeto por tipo de serviço necessário. Em vez de listar só tecnologias soltas, aqui fica claro qual ferramenta entra em cada parte do sistema e por quê.')

H('2. Serviço necessário para o RAG', 1)
P('A camada de RAG precisa resolver cinco coisas: receber arquivos, armazenar os arquivos, extrair/chunkar conteúdo, gerar embeddings e buscar os chunks relevantes.')
B('Backend do RAG: FastAPI')
B('Storage dos documentos: Cloudflare R2')
B('Banco vetorial: Qdrant')
B('Embeddings: OpenAI text-embedding-3-small (opção principal inicial)')
B('Banco relacional para metadata do app: Postgres (preferencialmente no Supabase)')

P('Onde cada um entra:')
B('FastAPI: recebe upload, dispara ingestão, expõe endpoints de search e chat. Importante: FastAPI é gratuita e open source. O custo não é da ferramenta em si; o custo é só da infraestrutura onde ela roda.', 0)
B('Cloudflare R2: guarda PDFs, imagens extraídas e artefatos do pipeline.', 0)
B('Qdrant: guarda vetores e metadata de retrieval.', 0)
B('OpenAI embeddings: transforma chunks em vetores numéricos.', 0)
B('Postgres: relaciona documentos, usuários, status de ingestão, permissões e sessões.', 0)

P('Alternativas para o RAG:')
B('Trocar Qdrant por pgvector se quiser reduzir o número de serviços.')
B('Trocar R2 por S3 ou Supabase Storage.')
B('Trocar OpenAI embeddings por outro provider, se custo/qualidade compensar.')

H('3. Serviço necessário para o chatbot', 1)
P('O chatbot precisa de interface, sessões, modelo de linguagem e conexão com o retrieval.')
B('Frontend do chat: Next.js')
B('Sessão do usuário: Supabase Auth')
B('Histórico e metadados da conversa: Postgres')
B('Orquestração do chat com RAG: FastAPI')
B('Modelo de linguagem: OpenAI, Anthropic ou Gemini')

P('Onde cada um entra:')
B('Next.js: tela do chat, UX, loading, streaming e histórico visual.', 0)
B('Supabase Auth: login e sessão do usuário.', 0)
B('Postgres: salva conversas, preferências, permissões e auditoria.', 0)
B('FastAPI: monta prompt, consulta retrieval, chama o modelo e devolve a resposta.', 0)

H('4. Serviço necessário para o site institucional', 1)
B('Next.js')
B('Vercel')
P('Aqui a ideia é simples: Next.js para construir landing page, páginas de produto, pricing, login e documentação pública. Vercel é a melhor opção de deploy se o frontend principal estiver em Next.js.')

H('5. Serviço necessário para login e sessão', 1)
B('Opção principal: Supabase Auth')
B('Alternativa: Auth.js (grátis e open source)')
P('Supabase Auth costuma ser mais fácil para MVP porque já resolve login, magic link, reset de senha e integração com Postgres. Auth.js é ótima alternativa se você quiser mais controle e menos dependência de fornecedor.')

H('6. Serviço necessário para o painel admin', 1)
B('Next.js para a interface')
B('Postgres para os dados administrativos')
B('FastAPI para endpoints operacionais e ações de ingestão/admin')
P('O painel admin deve centralizar usuários, documentos, status de ingestão, sessões, métricas, feedback de respostas e eventuais permissões por plano/perfil.')

H('7. Postgres no Supabase: não seria mais fácil?', 1)
P('Sim, no MVP geralmente é mais fácil.')
B('Recomendação prática: usar o Postgres do Supabase no começo.')
B('Motivo: já vem com auth, painel, banco, políticas e ecossistema integrado.')
B('Isso reduz bastante a fricção operacional no início.')
P('Ou seja: quando eu falo “Postgres”, na prática o caminho mais simples para você é “Supabase Postgres”, não um Postgres solto desde o dia 1.')
P('Resumo honesto:')
B('Se quer simplicidade: Supabase Auth + Supabase Postgres.')
B('Se quer mais independência e arquitetura desacoplada: Postgres separado.')

H('8. Comparação de modelos de IA por preço', 1)
P('Abaixo está uma comparação inicial de modelos úteis para o chatbot. Valores em USD por 1 milhão de tokens, com foco em referência pública de março/2026.')

t = doc.add_table(rows=1, cols=5)
t.style = 'Table Grid'
headers = ['Provider / Modelo', 'Input', 'Output', 'Perfil de uso', 'Observação']
for i, h in enumerate(headers):
    t.rows[0].cells[i].text = h
rows = [
    ['OpenAI GPT-5.4', '$2.50', '$15.00', 'Resposta de alta qualidade / casos complexos', 'Forte, mas caro para uso massivo'],
    ['OpenAI GPT-5.4 mini', '$0.75', '$4.50', 'Chat principal com bom custo-benefício', 'Muito interessante para MVP'],
    ['OpenAI GPT-5.4 nano', '$0.20', '$1.25', 'Tarefas simples / alto volume', 'Bom para rotinas leves'],
    ['Anthropic Claude Sonnet 4.6', '$3.00', '$15.00', 'Chat forte e raciocínio sólido', 'Faixa premium'],
    ['Anthropic Claude Haiku 4.5', '$1.00', '$5.00', 'Mais barato que Sonnet, mais leve', 'Boa opção intermediária'],
    ['Google Gemini 2.5 Pro', '$1.25', '$10.00', 'Bom equilíbrio custo/qualidade', 'Referência pública amplamente reportada'],
    ['Google Gemini Flash', '$0.30', '$2.50', 'Alto volume com custo menor', 'Boa opção econômica'],
    ['Google Gemini Flash-Lite', '$0.10', '$0.40', 'Tarefas muito baratas e simples', 'Extremamente econômico'],
]
for row in rows:
    c = t.add_row().cells
    for i, val in enumerate(row):
        c[i].text = val

P('Minha leitura prática:')
B('Se quiser melhor equilíbrio inicial: GPT-5.4 mini ou Gemini Flash.')
B('Se quiser máxima qualidade: GPT-5.4 ou Claude Sonnet 4.6.')
B('Se quiser esmagar custo: GPT-5.4 nano ou Gemini Flash-Lite.')

H('9. Comparação de embeddings', 1)
P('Para o RAG, o custo de embeddings também importa.')
B('OpenAI text-embedding-3-small: ~US$0.02 / 1M tokens (referência pública bem consolidada).')
B('Vantagem: barato, simples e muito usado.')
B('Desvantagem: ainda depende de provider externo.')
P('Para começar, ele continua sendo uma escolha bem razoável.')

H('10. Custos das ferramentas de infraestrutura', 1)

t2 = doc.add_table(rows=1, cols=4)
t2.style = 'Table Grid'
headers = ['Ferramenta', 'Preço de referência', 'Uso no projeto', 'Observação']
for i, h in enumerate(headers):
    t2.rows[0].cells[i].text = h
infra_rows = [
    ['FastAPI', 'Grátis', 'Backend de IA / RAG / chat', 'Open source; você só paga a hospedagem'],
    ['Supabase', 'Free; Pro a partir de US$25/mês', 'Auth + Postgres + recursos auxiliares', 'Provavelmente o jeito mais simples de começar'],
    ['Qdrant Cloud', 'Free tier + Standard por uso', 'Banco vetorial', 'Boa opção para RAG'],
    ['Cloudflare R2', 'US$0.015/GB-mês + operações', 'Storage de PDFs e artefatos', 'Bem competitivo'],
    ['Vercel', 'Hobby grátis; Pro US$20/usuário/mês + uso', 'Deploy do site e frontend', 'Excelente com Next.js'],
    ['Railway', 'Hobby US$5 mínimo; Pro US$20 mínimo + uso', 'Deploy do backend Python', 'MVP-friendly'],
    ['Auth.js', 'Grátis', 'Alternativa de autenticação', 'Open source; infra à parte'],
]
for row in infra_rows:
    c = t2.add_row().cells
    for i, val in enumerate(row):
        c[i].text = val

H('11. Stack recomendada final (mais simples e coerente)', 1)
B('RAG: FastAPI + Qdrant + Cloudflare R2 + OpenAI embeddings')
B('Chatbot: Next.js + FastAPI + modelo LLM')
B('Site: Next.js + Vercel')
B('Login e sessão: Supabase Auth')
B('Banco do app: Supabase Postgres')
B('Painel admin: Next.js + FastAPI + Supabase Postgres')

P('Ou seja: sim, faz bastante sentido usar o Postgres no Supabase no começo. Isso deixa o MVP mais simples sem te travar demais.')

H('12. Recomendação de escolha de modelo no MVP', 1)
P('Se eu fosse definir agora para primeira versão:')
B('LLM principal do chat: GPT-5.4 mini')
B('Embeddings: text-embedding-3-small')
B('Fallback ou comparação: Gemini Flash')
P('Essa combinação tende a equilibrar qualidade, simplicidade e custo.')

H('13. Fontes consultadas', 1)
B('OpenAI pricing: https://openai.com/api/pricing/')
B('Anthropic pricing: https://platform.claude.com/docs/en/about-claude/pricing')
B('Gemini pricing (referência pública): https://ai.google.dev/gemini-api/docs/pricing')
B('Qdrant pricing: https://qdrant.tech/pricing/')
B('Supabase pricing: https://supabase.com/pricing')
B('Vercel pricing: https://vercel.com/pricing')
B('Railway pricing: https://railway.com/pricing')
B('Cloudflare R2 pricing: https://developers.cloudflare.com/r2/pricing/')
B('Auth.js: https://authjs.dev/')

doc.save(out_path)
print(out_path)
