from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import RGBColor

out_path = r'C:\Users\vigna\.openclaw\workspace\RadioeXperienceRAG\catalog\Stack_MVP_RAG_Chatbot_Site_Admin.docx'

doc = Document()
styles = doc.styles
styles['Normal'].font.name = 'Calibri'
styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
styles['Normal'].font.size = Pt(11)

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Stack recomendado para o MVP\nRAG + Chatbot + Site + Login + Painel Admin')
r.bold = True
r.font.size = Pt(18)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Documento de decisão técnica e custos estimados — março/2026')
r.italic = True
r.font.size = Pt(10)


def heading(text, level=1):
    doc.add_heading(text, level=level)


def para(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    return p


def bullet(text, level=0):
    style = 'List Bullet' if level == 0 else 'List Bullet 2'
    doc.add_paragraph(text, style=style)

heading('1. Objetivo do projeto', 1)
para('Montar um produto inicial com base documental médica, chat com RAG, autenticação de usuários, site institucional e painel administrativo para gestão de conteúdo, sessões e operação.')

heading('2. Recomendação principal de stack', 1)
bullet('Frontend + painel admin: Next.js')
bullet('Autenticação: Supabase Auth (ou Auth.js como alternativa)')
bullet('Banco relacional do app: Postgres')
bullet('Backend de IA / RAG: FastAPI em Python')
bullet('Banco vetorial: Qdrant')
bullet('Armazenamento de arquivos: Cloudflare R2')
bullet('Deploy do frontend: Vercel')
bullet('Deploy do backend Python: Railway (ou Render)')

para('Resumo da lógica: Next.js cuida da experiência web e do painel; FastAPI cuida do que envolve ingestão, embeddings, retrieval e chat; Postgres guarda o estado da aplicação; Qdrant guarda os vetores; R2 guarda PDFs, imagens e assets do pipeline.')

heading('3. Ferramentas e onde cada uma entra', 1)

heading('3.1 Next.js', 2)
para('O que é: framework React full-stack para site, dashboard e rotas web.')
bullet('Onde usar: site público, área logada, painel admin, páginas de login, telas de chat e gestão de documentos.')
bullet('Por que faz sentido: produtividade alta, bom ecossistema, ótimo encaixe com painel/admin e frontend moderno.')
bullet('Alternativas: Remix, Nuxt, Astro + backend separado.')

heading('3.2 FastAPI', 2)
para('O que é: framework Python para APIs rápidas e tipadas.')
bullet('Onde usar: upload de documentos, pipeline de ingestão, chunking, embeddings, retrieval, endpoints do chat e tarefas assíncronas de indexação.')
bullet('Por que faz sentido: o núcleo do RAG vive melhor em Python; reduz atrito com parsing de PDF, embeddings e integração com Qdrant.')
bullet('Alternativas: Flask, Django Ninja, backend em Next.js + workers Python separados.')

heading('3.3 Postgres', 2)
para('O que é: banco relacional principal da aplicação.')
bullet('Onde usar: usuários, permissões, sessões de chat, documentos cadastrados, jobs, auditoria, configurações, billing futuro.')
bullet('Por que faz sentido: sólido, padrão de mercado, ótimo para app web com login e painel.')
bullet('Alternativas: MySQL, Neon Postgres, Supabase Postgres, Railway Postgres.')

heading('3.4 Qdrant', 2)
para('O que é: banco vetorial para embeddings e busca semântica.')
bullet('Onde usar: indexação dos chunks, metadata por documento, filtros por especialidade/tipo e recuperação semântica no chat.')
bullet('Por que faz sentido: foco bom em vetor search, API clara e ótimo fit para RAG.')
bullet('Alternativas: pgvector, Weaviate, Pinecone, Milvus.')

heading('3.5 Supabase Auth', 2)
para('O que é: camada pronta de autenticação e gestão básica de usuários.')
bullet('Onde usar: login, sessão, recuperação de senha, proteção da área logada e talvez banco Postgres se você quiser centralizar tudo lá.')
bullet('Por que faz sentido: acelera muito o MVP.')
bullet('Alternativas: Auth.js / NextAuth, Clerk, Auth0.')

heading('3.6 Cloudflare R2', 2)
para('O que é: object storage compatível com S3.')
bullet('Onde usar: PDFs brutos, imagens extraídas, versões processadas, arquivos de apoio do pipeline.')
bullet('Por que faz sentido: storage barato e sem taxa de egress direta do R2.')
bullet('Alternativas: AWS S3, Supabase Storage, Vercel Blob.')

heading('3.7 Vercel', 2)
para('O que é: plataforma de deploy ideal para Next.js.')
bullet('Onde usar: frontend, dashboard, landing e rotas leves do app web.')
bullet('Por que faz sentido: deploy simples, CI/CD fácil e ótimo DX com Next.js.')
bullet('Alternativas: Netlify, Cloudflare Pages, self-host em VPS.')

heading('3.8 Railway / Render', 2)
para('O que é: plataformas para subir backend Python, workers e serviços auxiliares.')
bullet('Onde usar: FastAPI, workers de ingestão, jobs assíncronos e eventualmente Postgres se não estiver no Supabase.')
bullet('Por que faz sentido: simplifica o deploy do backend sem virar um DevOps inteiro.')
bullet('Alternativas: Render, Fly.io, Hetzner/VPS próprio, AWS ECS.')

heading('4. Preços pesquisados (mar/2026)', 1)
para('Observação: preços mudam. Esses números são referências iniciais e devem ser confirmados na contratação.')

# Pricing table
rows = [
    ['Ferramenta', 'Preço base / referência', 'Observação'],
    ['OpenAI API', 'GPT-5.4: $2.50/1M input, $15/1M output; GPT-5.4 mini: $0.75/1M input, $4.50/1M output', 'Preço oficial da página de pricing da OpenAI'],
    ['OpenAI embeddings', 'text-embedding-3-small: ~US$0.02 / 1M tokens', 'Referência pública amplamente replicada; confirmar no painel oficial antes de fechar custo'],
    ['Qdrant Cloud', 'Free tier: 0.5 vCPU / 1GB RAM / 4GB disk; Standard: uso conforme recursos', 'Cobrança por compute, RAM, storage, backup e inference tokens'],
    ['Supabase', 'Free; Pro a partir de US$25/mês por projeto', 'Bom para auth + Postgres + storage básico'],
    ['Vercel', 'Hobby grátis; Pro US$20/usuário/mês + uso', 'Inclui US$20 de créditos mensais no Pro'],
    ['Railway', 'Free trial + depois US$1/mês no Free; Hobby US$5 mínimo; Pro US$20 mínimo', 'Uso cobrado por CPU, RAM, volume e egress'],
    ['Cloudflare R2', 'Storage US$0.015/GB-mês; Class A US$4.50/milhão; Class B US$0.36/milhão', '10 GB storage + 1M Class A + 10M Class B grátis/mês'],
    ['Auth.js', 'Grátis / open source', 'Sem custo da ferramenta; você paga a infra onde rodar'],
]

t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
for i, h in enumerate(rows[0]):
    t.rows[0].cells[i].text = h
for row in rows[1:]:
    c = t.add_row().cells
    for i, val in enumerate(row):
        c[i].text = val

heading('5. Alternativas por camada', 1)

heading('5.1 Se quiser simplificar o stack', 2)
bullet('Next.js + Supabase + Qdrant + R2 + workers Python separados')
bullet('Vantagem: menos peças visíveis no começo')
bullet('Desvantagem: parte do backend fica espalhada entre TS e Python')

heading('5.2 Se quiser tudo mais “web SaaS”', 2)
bullet('Next.js full-stack + Auth.js + Postgres + pgvector')
bullet('Vantagem: menos moving parts')
bullet('Desvantagem: pior separação entre app web e pipeline pesado de IA')

heading('5.3 Se quiser mais robustez enterprise', 2)
bullet('Next.js + FastAPI + Postgres gerenciado + Qdrant Cloud + S3 + fila (Redis/Celery)')
bullet('Vantagem: cresce melhor')
bullet('Desvantagem: mais custo e mais complexidade')

heading('6. Minha recomendação final para o MVP', 1)
para('Se a prioridade é colocar de pé rápido sem bagunçar o projeto, eu escolheria:')
bullet('Next.js para site + área logada + painel admin')
bullet('Supabase Auth para autenticação')
bullet('Postgres para dados transacionais')
bullet('FastAPI para o backend de IA/RAG')
bullet('Qdrant para vetor search')
bullet('Cloudflare R2 para arquivos')
bullet('Vercel no frontend')
bullet('Railway no backend Python')

para('Esse stack é equilibrado porque não é o mais “chique”, mas também não é gambiarra. Ele te deixa lançar, aprender com o uso e crescer depois sem refazer tudo do zero.')

heading('7. Custo inicial estimado', 1)
bullet('Modo econômico / protótipo: US$0 a US$50/mês + uso de API')
bullet('MVP mais sério: ~US$50 a US$150/mês + uso de OpenAI')
bullet('Maior variável no começo: consumo de LLM e embeddings')

heading('8. Fontes consultadas', 1)
bullet('OpenAI pricing: https://openai.com/api/pricing/')
bullet('Qdrant pricing: https://qdrant.tech/pricing/')
bullet('Supabase pricing: https://supabase.com/pricing')
bullet('Vercel pricing: https://vercel.com/pricing')
bullet('Railway pricing: https://railway.com/pricing')
bullet('Cloudflare R2 pricing: https://developers.cloudflare.com/r2/pricing/')
bullet('Auth.js: https://authjs.dev/')

doc.save(out_path)
print(out_path)
